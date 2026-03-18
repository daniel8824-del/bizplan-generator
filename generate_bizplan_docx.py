#!/usr/bin/env python3
"""
generate_bizplan_docx.py
========================
Generates a professional 예비창업패키지 사업계획서 DOCX document from a Markdown file.

Usage:
    python3 generate_bizplan_docx.py [input.md] [output.docx]

Defaults:
    input:  /mnt/c/Users/daniel/Desktop/사업계획서/workspace/사업계획서_최종_v2.md
    output: same directory as input, with .docx extension
"""

import sys
import os
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
DEFAULT_MD_PATH = (
    r"/mnt/c/Users/daniel/Desktop/사업계획서/workspace/사업계획서_최종_v2.md"
)
FONT_NAME = "Pretendard"
FONT_FALLBACK = "맑은 고딕"

BODY_SIZE = Pt(10)
SUB_BULLET_SIZE = Pt(10)
MAIN_BULLET_SIZE = Pt(11)
SECTION_TITLE_SIZE = Pt(13)
COVER_TITLE_SIZE = Pt(22)

MARGIN_TOP = Cm(2)
MARGIN_BOTTOM = Cm(2)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(2.5)

# Perplexity 스타일 색상
TEAL = "01696F"
TEAL_LIGHT = "E0F2F1"
SECTION_SHADING_COLOR = TEAL
TABLE_HEADER_COLOR = TEAL
TABLE_ALT_ROW_COLOR = "F5F5F5"
BORDER_COLOR = "CCCCCC"

# Pattern to detect numbered section titles like "1. 문제 인식" or "Ⅰ. 개요"
SECTION_TITLE_RE = re.compile(
    r"^(#{1,2})\s+"
    r"((?:\d+\.|[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]+\.?)\s+.+)$"
)

# Pattern for □ sections like "□ 일반현황"
CHECKBOX_SECTION_RE = re.compile(r"^(#{1,3})\s+(□\s+.+)$")

# Pattern for markdown table separator row
TABLE_SEP_RE = re.compile(r"^\s*\|[\s\-:|]+\|[\s\-:|]*$")

# Bold pattern **text**
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

# Image pattern ![alt](path)
IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


# ---------------------------------------------------------------------------
# Helper: set font for a run
# ---------------------------------------------------------------------------
def set_run_font(run, font_name=FONT_NAME, size=BODY_SIZE, bold=False, color=None):
    """Configure font properties for a run, including East Asian font."""
    run.font.size = size
    run.font.bold = bold
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    # Set East Asian font (critical for Korean text rendering)
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


# ---------------------------------------------------------------------------
# Helper: add shaded bordered box around a paragraph (section title style)
# ---------------------------------------------------------------------------
def add_section_title_shading(paragraph, shade_color=SECTION_SHADING_COLOR):
    """Add teal background shading with white text for section titles."""
    pPr = paragraph._element.get_or_add_pPr()

    # Teal 배경
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{shade_color}" w:val="clear"/>'
    )
    pPr.append(shading)

    # 둥근 느낌의 테두리 (teal 색상)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="2" w:color="{shade_color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="6" w:color="{shade_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="2" w:color="{shade_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="6" w:color="{shade_color}"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)

    # 텍스트를 흰색으로
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Helper: set table cell shading
# ---------------------------------------------------------------------------
def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    tcPr.append(shading)


# ---------------------------------------------------------------------------
# Helper: set thin borders on a table
# ---------------------------------------------------------------------------
def set_table_borders(table, color=BORDER_COLOR, sz="4"):
    """Apply thin borders to all cells of a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f"<w:tblPr {nsdecls('w')}/>"
    )
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f"</w:tblBorders>"
    )
    borders = parse_xml(borders_xml)
    # Remove existing borders if present
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


# ---------------------------------------------------------------------------
# Helper: add text with inline bold parsing
# ---------------------------------------------------------------------------
def add_mixed_text(paragraph, text, font_name=FONT_NAME, size=BODY_SIZE,
                   base_bold=False, color=None):
    """Add text to a paragraph, converting **bold** markers to bold runs."""
    parts = BOLD_RE.split(text)
    for i, part in enumerate(parts):
        if not part:
            continue
        is_bold_segment = (i % 2 == 1)  # odd indices are captured bold groups
        run = paragraph.add_run(part)
        set_run_font(
            run,
            font_name=font_name,
            size=size,
            bold=base_bold or is_bold_segment,
            color=color,
        )


# ---------------------------------------------------------------------------
# Helper: add page number footer  "- N -"
# ---------------------------------------------------------------------------
def add_page_number_footer(section):
    """Add centered page number in '- N -' format to the section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = p.add_run("- ")
    set_run_font(run1, size=Pt(9))

    # PAGE field
    fld_xml = (
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    )
    run_fld_begin = p.add_run()
    run_fld_begin._element.append(parse_xml(fld_xml))

    run_instr = p.add_run()
    instr = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
    )
    run_instr._element.append(instr)

    run_fld_end = p.add_run()
    run_fld_end._element.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    )

    run2 = p.add_run(" -")
    set_run_font(run2, size=Pt(9))


# ---------------------------------------------------------------------------
# Parse markdown into structured blocks
# ---------------------------------------------------------------------------
def parse_markdown(md_text):
    """
    Parse markdown text into a list of block dicts.
    Block types:
        - 'heading': level, text
        - 'section_title': text (numbered section like "1. 문제 인식")
        - 'checkbox_section': text (□ sections)
        - 'table': rows (list of lists of cell text)
        - 'bullet_main': text (◦ bullet)
        - 'bullet_sub': text (- bullet under ◦)
        - 'paragraph': text
        - 'page_break'
        - 'blank'
    """
    lines = md_text.split("\n")
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            blocks.append({"type": "blank"})
            i += 1
            continue

        # Image ![alt](path)
        m_img = IMAGE_RE.match(stripped)
        if m_img:
            blocks.append({
                "type": "image",
                "alt": m_img.group(1),
                "path": m_img.group(2),
            })
            i += 1
            continue

        # Page break (horizontal rule or explicit)
        if stripped in ("---", "***", "___", "<!-- pagebreak -->"):
            blocks.append({"type": "page_break"})
            i += 1
            continue

        # Heading with checkbox section  □
        m_cb = CHECKBOX_SECTION_RE.match(stripped)
        if m_cb:
            blocks.append({
                "type": "checkbox_section",
                "level": len(m_cb.group(1)),
                "text": m_cb.group(2),
            })
            i += 1
            continue

        # Heading with numbered section title
        m_sec = SECTION_TITLE_RE.match(stripped)
        if m_sec:
            blocks.append({
                "type": "section_title",
                "level": len(m_sec.group(1)),
                "text": m_sec.group(2),
            })
            i += 1
            continue

        # Generic heading
        m_head = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if m_head:
            blocks.append({
                "type": "heading",
                "level": len(m_head.group(1)),
                "text": m_head.group(2),
            })
            i += 1
            continue

        # Table (collect contiguous table lines)
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            # Parse table
            rows = []
            for tl in table_lines:
                if TABLE_SEP_RE.match(tl):
                    continue  # skip separator row
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        # ◦ bullet (main bullet)
        if stripped.startswith("◦") or stripped.startswith("◦"):
            text = stripped.lstrip("◦").strip()
            blocks.append({"type": "bullet_main", "text": text})
            i += 1
            continue

        # - bullet (sub-bullet / dash item)
        if re.match(r"^\s*[-]\s+", line):
            text = re.sub(r"^\s*[-]\s+", "", line).strip()
            blocks.append({"type": "bullet_sub", "text": text})
            i += 1
            continue

        # * or • bullet (treat similarly to ◦)
        if stripped.startswith("•") or re.match(r"^\s*\*\s+", stripped):
            text = re.sub(r"^[•\*]\s*", "", stripped).strip()
            blocks.append({"type": "bullet_main", "text": text})
            i += 1
            continue

        # Regular paragraph
        blocks.append({"type": "paragraph", "text": stripped})
        i += 1

    return blocks


# ---------------------------------------------------------------------------
# Build DOCX from parsed blocks
# ---------------------------------------------------------------------------
def build_docx(blocks, title="예비창업패키지 예비창업자 사업계획서", base_dir=None):
    """Create a Document from parsed blocks and return it."""
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # Footer with page numbers
    add_page_number_footer(section)

    # -- Default style adjustments --
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.line_spacing = 1.15
    # Set East Asian font on the default style
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        style.element.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)

    # -- Cover page (Perplexity 스타일) --
    # Teal 상단 라인
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(0)
    pPr = p_line._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="24" w:space="1" w:color="{TEAL}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)

    # 타이틀
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(8)
    run_title = p_title.add_run(title)
    set_run_font(run_title, size=COVER_TITLE_SIZE, bold=True,
                 color=RGBColor(0x01, 0x69, 0x6F))

    # 부제
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sub.paragraph_format.space_before = Pt(4)
    run_sub = p_sub.add_run("예비창업패키지 대비 사업계획서")
    set_run_font(run_sub, size=Pt(13), bold=False,
                 color=RGBColor(0x66, 0x66, 0x66))

    for _ in range(3):
        doc.add_paragraph()

    # 하단 구분선
    p_line2 = doc.add_paragraph()
    pPr2 = p_line2._element.get_or_add_pPr()
    pBdr2 = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="1" w:color="{TEAL}"/>'
        f'</w:pBdr>'
    )
    pPr2.append(pBdr2)

    # 작성 정보
    p_info = doc.add_paragraph()
    p_info.paragraph_format.space_before = Pt(8)
    run_info = p_info.add_run(f"작성 기준일: 2026년 3월")
    set_run_font(run_info, size=Pt(10), color=RGBColor(0x66, 0x66, 0x66))

    doc.add_page_break()

    # -- Render blocks --
    # Skip the first H1 heading if it matches the cover title (avoid blank page 2)
    skip_first_h1 = True
    is_first_section = True
    for block in blocks:
        btype = block["type"]

        if btype == "blank":
            continue

        elif btype == "page_break":
            doc.add_page_break()

        elif btype == "section_title":
            # Major numbered section heading with shaded box
            if not is_first_section:
                # Add some spacing (not a full page break unless explicit)
                p_spacer = doc.add_paragraph()
                p_spacer.paragraph_format.space_after = Pt(6)
            is_first_section = False

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(8)
            add_mixed_text(p, block["text"], size=SECTION_TITLE_SIZE, base_bold=True)
            add_section_title_shading(p)

        elif btype == "checkbox_section":
            # □ sections
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            add_mixed_text(p, block["text"], size=Pt(11), base_bold=True)

        elif btype == "heading":
            level = block["level"]
            # Skip first H1 — it duplicates the cover page title
            if skip_first_h1 and level == 1:
                skip_first_h1 = False
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            if level <= 2:
                add_mixed_text(p, block["text"], size=Pt(12), base_bold=True)
                if level == 2:
                    add_section_title_shading(p)
            elif level == 3:
                add_mixed_text(p, block["text"], size=Pt(11), base_bold=True)
            else:
                add_mixed_text(p, block["text"], size=BODY_SIZE, base_bold=True)

        elif btype == "table":
            rows = block["rows"]
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=0, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(table)

            # Calculate available width
            avail_width = Cm(21) - MARGIN_LEFT - MARGIN_RIGHT

            for row_idx, row_data in enumerate(rows):
                row = table.add_row()
                for col_idx in range(num_cols):
                    cell = row.cells[col_idx]
                    cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                    p = cell.paragraphs[0]
                    p.clear()

                    if row_idx == 0:
                        # 헤더: Teal 배경 + 흰색 텍스트
                        add_mixed_text(p, cell_text, size=Pt(9), base_bold=True,
                                      color=RGBColor(0xFF, 0xFF, 0xFF))
                        set_cell_shading(cell, TABLE_HEADER_COLOR)
                    else:
                        add_mixed_text(p, cell_text, size=Pt(9), base_bold=False)
                        # 교차 배경색
                        if row_idx % 2 == 0:
                            set_cell_shading(cell, TABLE_ALT_ROW_COLOR)

                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)

            # Add small spacing after table
            p_after = doc.add_paragraph()
            p_after.paragraph_format.space_after = Pt(4)

        elif btype == "bullet_main":
            # ◦ main bullet — bold, 11pt
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5)

            run_marker = p.add_run("◦ ")
            set_run_font(run_marker, size=MAIN_BULLET_SIZE, bold=True,
                        color=RGBColor(0x01, 0x69, 0x6F))
            add_mixed_text(p, block["text"], size=MAIN_BULLET_SIZE, base_bold=True)

        elif btype == "bullet_sub":
            # - sub bullet — regular, 10pt, indented
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(1.2)
            p.paragraph_format.first_line_indent = Cm(-0.4)

            run_marker = p.add_run("- ")
            set_run_font(run_marker, size=SUB_BULLET_SIZE, bold=False)
            add_mixed_text(p, block["text"], size=SUB_BULLET_SIZE)

        elif btype == "image":
            # Insert image from file path
            img_path = block["path"]
            if base_dir and not os.path.isabs(img_path):
                img_path = os.path.join(base_dir, img_path)
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run()
                # Calculate proportional size preserving aspect ratio
                try:
                    from PIL import Image as PILImage
                    pil_img = PILImage.open(img_path)
                    img_w, img_h = pil_img.size
                    max_width = Cm(15)
                    max_height = Cm(10)
                    aspect = img_w / img_h
                    width = max_width
                    height = Emu(int(width / aspect))
                    if height > max_height:
                        height = max_height
                        width = Emu(int(height * aspect))
                    run.add_picture(img_path, width=width, height=height)
                except (ImportError, Exception):
                    run.add_picture(img_path, width=Cm(15))
                # Add caption below if alt text exists
                if block.get("alt"):
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_before = Pt(1)
                    p_cap.paragraph_format.space_after = Pt(6)
                    run_cap = p_cap.add_run(block["alt"])
                    set_run_font(run_cap, size=Pt(8), bold=False,
                                color=RGBColor(0x66, 0x66, 0x66))

        elif btype == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_mixed_text(p, block["text"], size=BODY_SIZE)

    return doc


# ---------------------------------------------------------------------------
# Convert DOCX to PDF via LibreOffice
# ---------------------------------------------------------------------------
def convert_to_pdf(docx_path):
    """
    Convert a DOCX file to PDF using LibreOffice in headless mode.
    Returns the expected PDF path, or None if conversion failed.
    """
    docx_path = os.path.abspath(docx_path)
    output_dir = os.path.dirname(docx_path)
    pdf_name = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
    pdf_path = os.path.join(output_dir, pdf_name)

    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", docx_path],
            cwd=output_dir,
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode == 0 and os.path.exists(pdf_path):
            print(f"PDF generated: {pdf_path}")
            return pdf_path
        else:
            print(f"PDF conversion warning: {result.stderr.strip()}")
            if os.path.exists(pdf_path):
                return pdf_path
            return None
    except FileNotFoundError:
        print("LibreOffice (soffice) not found. Skipping PDF conversion.")
        print("Install with: sudo apt install libreoffice")
        return None
    except subprocess.TimeoutExpired:
        print("PDF conversion timed out.")
        return None


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    # Determine input/output paths
    if len(sys.argv) >= 2:
        md_path = sys.argv[1]
    else:
        md_path = DEFAULT_MD_PATH

    if len(sys.argv) >= 3:
        docx_path = sys.argv[2]
    else:
        docx_path = os.path.splitext(md_path)[0] + ".docx"

    # Ensure output directory exists
    output_dir = os.path.dirname(docx_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    # Read markdown
    if not os.path.exists(md_path):
        print(f"Error: Markdown file not found: {md_path}")
        print("Usage: python3 generate_bizplan_docx.py [input.md] [output.docx]")
        sys.exit(1)

    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    print(f"Reading: {md_path}")

    # Parse and build
    blocks = parse_markdown(md_text)
    print(f"Parsed {len(blocks)} blocks from markdown.")

    base_dir = os.path.dirname(os.path.abspath(md_path))
    doc = build_docx(blocks, base_dir=base_dir)

    # Save
    doc.save(docx_path)
    print(f"DOCX generated: {docx_path}")

    # Attempt PDF conversion
    convert_to_pdf(docx_path)

    return docx_path


if __name__ == "__main__":
    main()
