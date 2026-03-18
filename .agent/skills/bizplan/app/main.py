"""
bizplan-app — 예비창업패키지 사업계획서 자동 생성 웹앱
"""

from fastapi import FastAPI, Request, Form, BackgroundTasks, UploadFile, File
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import asyncio
import os
import uuid
from datetime import datetime
from typing import Optional

from app.research import run_all_research, firecrawl_search
from app.file_parser import extract_text
from app.writer import (
    write_phase1, write_market_analysis, write_target_customers,
    write_competitors, write_go_to_market, write_founder_fit,
    write_final_bizplan, call_openrouter
)

app = FastAPI(title="BizPlan Generator", version="1.0")

templates = Jinja2Templates(directory="app/templates")
app.mount("/static", StaticFiles(directory="app/static"), name="static")

# 진행 상태 저장 (메모리)
jobs: dict = {}

WORKSPACE = os.environ.get("WORKSPACE", "/tmp/bizplan-workspace")


@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/api/start")
async def start_bizplan(
    background_tasks: BackgroundTasks,
    item: str = Form(...),
    problem: str = Form(""),
    target: str = Form(""),
    stage: str = Form("아이디어"),
    founder: str = Form(""),
    team: str = Form(""),
    company: str = Form(""),
    openrouter_key: str = Form(None),
    upload_file: list[UploadFile] = File([]),
    upload_intent: str = Form(""),
):
    """사업계획서 생성 시작"""
    job_id = str(uuid.uuid4())[:8]
    project_dir = os.path.join(WORKSPACE, f"{company}_{datetime.now().strftime('%Y%m%d_%H%M')}")
    os.makedirs(os.path.join(project_dir, "charts"), exist_ok=True)

    # 파일 업로드 처리 (다중 파일)
    uploaded_text = ""
    for uf in upload_file:
        if uf and uf.filename:
            upload_dir = os.path.join(project_dir, "uploads")
            os.makedirs(upload_dir, exist_ok=True)
            file_path = os.path.join(upload_dir, uf.filename)
            with open(file_path, "wb") as f:
                content = await uf.read()
                f.write(content)
            text = extract_text(file_path) or ""
            if text:
                uploaded_text += f"\n\n--- {uf.filename} ---\n{text}"
            print(f"[Upload] {uf.filename} ({len(text)} chars), intent={upload_intent}")

    jobs[job_id] = {
        "status": "started",
        "phase": "Phase 1: 아이템 구체화",
        "progress": 0,
        "project_dir": project_dir,
        "results": {},
    }

    item_info = {
        "item": item, "problem": problem, "target": target,
        "stage": stage, "founder": founder, "team": team, "company": company,
        "uploaded_text": uploaded_text,
        "upload_intent": upload_intent,
    }

    api_key = os.environ.get("OPENROUTER_API_KEY") or openrouter_key
    if not api_key:
        return JSONResponse({"error": "API 키가 설정되지 않았습니다."}, status_code=400)

    background_tasks.add_task(run_pipeline, job_id, item_info, api_key)

    return {"job_id": job_id, "status": "started"}


@app.get("/api/status/{job_id}")
async def get_status(job_id: str):
    """진행 상태 조회"""
    job = jobs.get(job_id)
    if not job:
        return JSONResponse({"error": "Job not found"}, status_code=404)
    return {
        "status": job["status"],
        "phase": job["phase"],
        "progress": job["progress"],
        "results": {k: v[:200] + "..." if isinstance(v, str) and len(v) > 200 else v
                    for k, v in job.get("results", {}).items()
                    if k != "final_bizplan"},
    }


@app.get("/api/download/{job_id}/{filetype}")
async def download_file(job_id: str, filetype: str):
    """결과 파일 다운로드 (md, pdf)"""
    job = jobs.get(job_id)
    if not job:
        return JSONResponse({"error": "Job not found"}, status_code=404)

    project_dir = job["project_dir"]

    if filetype == "md":
        path = os.path.join(project_dir, "06_사업계획서_최종.md")
    elif filetype == "pdf":
        path = os.path.join(project_dir, "사업계획서.pdf")
    elif filetype == "docx":
        path = os.path.join(project_dir, "사업계획서.docx")
    else:
        return JSONResponse({"error": "Invalid filetype"}, status_code=400)

    if os.path.exists(path):
        return FileResponse(path, filename=os.path.basename(path))
    return JSONResponse({"error": "File not found"}, status_code=404)


@app.get("/api/files/{job_id}")
async def list_files(job_id: str):
    """프로젝트 디렉토리의 모든 파일 목록"""
    job = jobs.get(job_id)
    if not job:
        return JSONResponse({"error": "Job not found"}, status_code=404)

    project_dir = job["project_dir"]
    if not os.path.isdir(project_dir):
        return JSONResponse({"error": "Directory not found"}, status_code=404)

    files = []
    for fname in sorted(os.listdir(project_dir)):
        fpath = os.path.join(project_dir, fname)
        if os.path.isfile(fpath) and not fname.startswith('.'):
            size = os.path.getsize(fpath)
            files.append({"name": fname, "size": size})

    # charts/ subdirectory
    charts_dir = os.path.join(project_dir, "charts")
    if os.path.isdir(charts_dir):
        for fname in sorted(os.listdir(charts_dir)):
            fpath = os.path.join(charts_dir, fname)
            if os.path.isfile(fpath):
                files.append({"name": f"charts/{fname}", "size": os.path.getsize(fpath)})

    return {"files": files}


@app.get("/api/download/{job_id}/file/{filename:path}")
async def download_generic(job_id: str, filename: str):
    """개별 파일 다운로드"""
    if ".." in filename:
        return JSONResponse({"error": "Invalid"}, status_code=400)
    job = jobs.get(job_id)
    if not job:
        return JSONResponse({"error": "Job not found"}, status_code=404)

    file_path = os.path.join(job["project_dir"], filename)
    if not os.path.isfile(file_path):
        return JSONResponse({"error": "File not found"}, status_code=404)

    real_path = os.path.realpath(file_path)
    real_project = os.path.realpath(job["project_dir"])
    if not real_path.startswith(real_project):
        return JSONResponse({"error": "Access denied"}, status_code=403)

    return FileResponse(file_path, filename=os.path.basename(file_path))


async def run_pipeline(job_id: str, item_info: dict, api_key: str):
    """전체 파이프라인 백그라운드 실행"""
    job = jobs[job_id]
    project_dir = job["project_dir"]

    try:
        # 업로드 파일 컨텍스트 구성
        upload_context = ""
        if item_info.get("uploaded_text"):
            intent = item_info.get("upload_intent", "")
            upload_context = f"""

## 사용자 업로드 자료
> 사용자 의도: {intent if intent else '참고자료로 활용'}

{item_info['uploaded_text'][:8000]}
"""

        # Phase 1: 아이템 구체화
        job["phase"] = "Phase 1: 아이템 구체화"
        job["progress"] = 10
        # 업로드 자료가 있으면 item_info에 컨텍스트 추가
        if upload_context:
            item_info["_upload_context"] = upload_context
        phase1 = await write_phase1(item_info, api_key)
        _save(project_dir, "00_item_brief.md", phase1)
        job["results"]["phase1"] = "완료"
        job["progress"] = 20

        # Phase 2: 리서치 (병렬)
        job["phase"] = "Phase 2: 시장 조사 (병렬 실행 중)"
        job["progress"] = 25

        industry = item_info["item"]
        target = item_info.get("target", "")

        research = await run_all_research(industry, target)

        # 시장 규모 분석
        job["phase"] = "Phase 2-1: 시장 규모 분석"
        job["progress"] = 35
        market_report = await write_market_analysis(phase1, research["market_size"], api_key)
        _save(project_dir, "01_market_size.md", market_report)

        # 타겟 고객
        job["phase"] = "Phase 2-2: 타겟 고객 세분화"
        job["progress"] = 45
        target_report = await write_target_customers(phase1, research["target_customers"], api_key)
        _save(project_dir, "02_target_customers.md", target_report)

        # 경쟁사
        job["phase"] = "Phase 2-3: 경쟁사 분석"
        job["progress"] = 55
        comp_report = await write_competitors(phase1, research["competitors"], api_key)
        _save(project_dir, "03_competitors.md", comp_report)

        # 진입 전략
        job["phase"] = "Phase 2-4: 시장 진입 전략"
        job["progress"] = 62
        gtm_report = await write_go_to_market(phase1, research["go_to_market"], api_key)
        _save(project_dir, "04_go_to_market.md", gtm_report)

        # Founder-Market Fit
        job["phase"] = "Phase 2-5: Founder-Market Fit 분석"
        job["progress"] = 68
        founder_report = await write_founder_fit(phase1, item_info, api_key)
        _save(project_dir, "05_founder_fit.md", founder_report)

        job["results"]["research"] = "5개 리서치 완료 (시장+타겟+경쟁사+진입+Founder Fit)"
        job["progress"] = 72

        # 차트 생성 — AI가 데이터를 보고 필요한 차트를 판단하여 생성
        job["phase"] = "차트 생성 중 (AI 판단)"
        job["progress"] = 74
        await _generate_charts_ai(project_dir, market_report, comp_report, target_report, api_key)

        # Phase 3: 사업계획서 작성
        job["phase"] = "Phase 3: 사업계획서 작성 중"
        job["progress"] = 78

        all_research = f"""# 아이템 브리프
{phase1}
{upload_context}
# 시장 규모 분석
{market_report}

# 타겟 고객 세분화
{target_report}

# 경쟁사 분석
{comp_report}

# 시장 진입 전략
{gtm_report}

# Founder-Market Fit
{founder_report}"""

        bizplan = await write_final_bizplan(all_research, item_info, api_key)
        _save(project_dir, "06_사업계획서_최종.md", bizplan)
        job["results"]["bizplan"] = "작성 완료"
        job["progress"] = 90

        # Phase 4: DOCX + PDF 생성
        job["phase"] = "Phase 4: DOCX + PDF 생성 중"
        job["progress"] = 92
        import subprocess

        md_path = os.path.join(project_dir, "06_사업계획서_최종.md")
        docx_path = os.path.join(project_dir, "사업계획서.docx")
        pdf_path = os.path.join(project_dir, "사업계획서.pdf")

        # 스크립트 경로: /app/ (Railway) 또는 ~/bizplan-app/ (로컬)
        app_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        docx_script = os.path.join(app_root, "generate_bizplan_docx.py")
        pdf_script = os.path.join(app_root, "generate_bizplan_pdf.py")

        print(f"[Phase 4] md: {md_path}")
        print(f"[Phase 4] docx_script: {docx_script} exists={os.path.exists(docx_script)}")
        print(f"[Phase 4] pdf_script: {pdf_script} exists={os.path.exists(pdf_script)}")

        # DOCX 생성 + 차트 삽입
        try:
            if os.path.exists(docx_script):
                result = subprocess.run(
                    ["python3", docx_script, md_path, docx_path],
                    capture_output=True, text=True, timeout=60
                )
                print(f"[DOCX] stdout: {result.stdout[:200]}")
                print(f"[DOCX] stderr: {result.stderr[:200]}")
                if os.path.exists(docx_path):
                    _insert_charts_docx(docx_path, os.path.join(project_dir, "charts"))
                    job["results"]["docx"] = "DOCX 생성 완료"
                else:
                    job["results"]["docx"] = "DOCX 파일 미생성"
            else:
                job["results"]["docx"] = f"DOCX 스크립트 없음: {docx_script}"
        except Exception as e:
            job["results"]["docx"] = f"DOCX 실패: {str(e)[:100]}"

        # PDF 생성
        job["phase"] = "Phase 4: PDF 생성 중"
        job["progress"] = 96
        try:
            if os.path.exists(pdf_script):
                result = subprocess.run(
                    ["python3", pdf_script, md_path, pdf_path],
                    capture_output=True, text=True, timeout=60
                )
                print(f"[PDF] stdout: {result.stdout[:200]}")
                print(f"[PDF] stderr: {result.stderr[:200]}")
                if os.path.exists(pdf_path):
                    job["results"]["pdf"] = "PDF 생성 완료"
                else:
                    job["results"]["pdf"] = "PDF 파일 미생성"
            else:
                job["results"]["pdf"] = f"PDF 스크립트 없음: {pdf_script}"
        except Exception as e:
            job["results"]["pdf"] = f"PDF 실패: {str(e)[:100]}"

        job["status"] = "completed"
        job["phase"] = "완료"
        job["progress"] = 100

    except Exception as e:
        job["status"] = "error"
        job["phase"] = f"에러: {str(e)[:200]}"


def _save(project_dir: str, filename: str, content: str):
    path = os.path.join(project_dir, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


async def _generate_charts_ai(project_dir: str, market_data: str, comp_data: str,
                               target_data: str, api_key: str):
    """AI가 리서치 데이터를 분석하여 필요한 차트를 판단하고 matplotlib 코드를 생성·실행"""
    charts_dir = os.path.join(project_dir, 'charts')
    os.makedirs(charts_dir, exist_ok=True)

    prompt = f"""아래 시장 조사 데이터를 분석하여, 사업계획서에 들어갈 차트를 생성하는 Python matplotlib 코드를 작성해주세요.

## 시장 규모 데이터
{market_data[:2000]}

## 경쟁사 데이터
{comp_data[:2000]}

## 타겟 고객 데이터
{target_data[:1000]}

## 요구사항
1. 데이터에서 실제 수치를 추출하여 차트에 반영
2. 차트 종류와 개수는 데이터를 보고 판단 (보통 2~4개)
3. 가능한 차트 유형:
   - TAM/SAM/SOM 시장 규모 바 차트
   - 시장 성장 추이 라인 차트
   - 경쟁사 포지셔닝 산점도 맵 (X: 자동화 수준, Y: 산업 특화도)
   - 경쟁사 기능 비교 히트맵
   - 시장 세그먼트 파이 차트
   - 고객 세그먼트 비교 바 차트
   - 가격대 비교 차트

## Python 코드 규칙
- `import matplotlib; matplotlib.use('Agg')` 필수
- 한글 폰트 설정 (반드시 이 코드를 맨 앞에 넣으세요):
```python
import matplotlib.font_manager as fm
import os
font_dir = os.path.expanduser('~/.fonts')
if os.path.exists(font_dir) and any('Pretendard' in f for f in os.listdir(font_dir)):
    for otf in os.listdir(font_dir):
        if otf.endswith('.otf') and 'Pretendard' in otf:
            fm.fontManager.addfont(os.path.join(font_dir, otf))
    plt.rcParams['font.family'] = 'Pretendard'
else:
    for f in fm.fontManager.ttflist:
        if 'Noto Sans CJK' in f.name:
            plt.rcParams['font.family'] = f.name
            break
    else:
        plt.rcParams['font.family'] = 'sans-serif'
        plt.rcParams['font.sans-serif'] = ['Noto Sans CJK KR', 'Noto Sans CJK', 'sans-serif']
```
- `plt.rcParams['axes.unicode_minus'] = False`
- `plt.rcParams['font.size'] = 12`
- figsize는 (10, 6) 이상으로 설정 (작은 figure에 텍스트가 겹치지 않도록)
- 폰트 크기 설정: `plt.rcParams['font.size'] = 12` (기본), 제목은 14~16pt
- 저장: `plt.savefig('{charts_dir}/[파일명].png', dpi=200, bbox_inches='tight', facecolor='white')`
- 각 차트 뒤에 `plt.close()` 필수
- 반드시 `plt.tight_layout()` 호출 후 savefig
- 차트 색상: Teal(#01696F) 메인, #E74C3C 강조, #F39C12 보조
- 우리 제품은 빨간 별(★) 마커로 강조
- 차트 제목과 레이블은 한국어
- 레이블이 겹치지 않도록 적절한 여백과 폰트 크기 조절

Python 코드만 출력하세요. 설명 없이 코드 블록만."""

    try:
        chart_code = await call_openrouter(prompt, api_key=api_key, max_tokens=4000)

        # 코드 블록 추출
        import re
        code_match = re.search(r'```python\s*(.*?)```', chart_code, re.DOTALL)
        if code_match:
            code = code_match.group(1)
        else:
            code = chart_code

        # charts_dir 변수를 코드에 주입
        code = f"charts_dir = '{charts_dir}'\n" + code

        # 실행
        exec_globals = {}
        exec(code, exec_globals)

        generated = [f for f in os.listdir(charts_dir) if f.endswith('.png')]
        print(f"AI 차트 생성 완료: {generated}")

    except Exception as e:
        print(f"AI 차트 생성 실패: {e}")
        # fallback: 기본 차트 생성
        _generate_charts_fallback(project_dir)


def _generate_charts_fallback(project_dir: str):
    """AI 차트 생성 실패 시 기본 차트 생성"""
    _generate_charts(project_dir, "", "")


def _insert_charts_docx(docx_path: str, charts_dir: str):
    """DOCX에 차트 이미지 삽입"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(docx_path)
        # 차트 디렉토리의 모든 PNG를 키워드 매칭으로 삽입
        keyword_map = {
            'market': ('시장 규모', '시장 규모 분석'),
            'trend': ('CAGR', '시장 성장 추이'),
            'competitor': ('경쟁사', '경쟁사 포지셔닝 맵'),
            'price': ('가격', '가격대 비교'),
            'segment': ('고객', '고객 세그먼트'),
            'revenue': ('수익', '수익 모델'),
            'roadmap': ('로드맵', '사업 로드맵'),
        }
        chart_files = {}
        for f in os.listdir(charts_dir):
            if not f.endswith('.png'):
                continue
            bn = f.replace('.png', '').lower()
            matched = False
            for key, (keyword, caption) in keyword_map.items():
                if key in bn:
                    chart_files[f] = (keyword, caption)
                    matched = True
                    break
            if not matched:
                # 파일명 자체를 키워드로 사용
                chart_files[f] = (bn.replace('_', ' '), bn.replace('_', ' ').title())

        for chart_file, (keyword, caption) in chart_files.items():
            chart_path = os.path.join(charts_dir, chart_file)
            if not os.path.exists(chart_path):
                continue
            # 키워드가 포함된 문단 뒤에 삽입
            for i, p in enumerate(doc.paragraphs):
                if keyword in p.text and chart_file.split('.')[0] not in str(getattr(p, '_inserted', '')):
                    new_p = doc.add_paragraph()
                    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = new_p.add_run()
                    # Proportional sizing preserving aspect ratio
                    try:
                        from PIL import Image as PILImage
                        pil_img = PILImage.open(chart_path)
                        img_w, img_h = pil_img.size
                        max_width = Cm(15)
                        max_height = Cm(10)
                        aspect = img_w / img_h
                        width = max_width
                        height = Emu(int(width / aspect))
                        if height > max_height:
                            height = max_height
                            width = Emu(int(height * aspect))
                        run.add_picture(chart_path, width=width, height=height)
                    except (ImportError, Exception):
                        run.add_picture(chart_path, width=Inches(5.2))
                    # 캡션
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = cap.add_run(f'< {caption} >')
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    r.font.italic = True
                    p._element.addnext(cap._element)
                    p._element.addnext(new_p._element)
                    break

        doc.save(docx_path)
    except Exception as e:
        print(f"DOCX 차트 삽입 실패: {e}")


def _generate_charts(project_dir: str, market_data: str, comp_data: str):
    """matplotlib 차트 생성"""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.font_manager as fm

        # 한글 폰트 설정 (Railway: Noto Sans CJK, 로컬: Pretendard)
        font_dir = os.path.expanduser('~/.fonts')
        if os.path.exists(font_dir):
            for otf in os.listdir(font_dir):
                if otf.endswith('.otf') and 'Pretendard' in otf:
                    fm.fontManager.addfont(os.path.join(font_dir, otf))
            plt.rcParams['font.family'] = 'Pretendard'
        else:
            # Railway 환경: Noto Sans CJK 사용
            for f in fm.fontManager.ttflist:
                if 'Noto Sans CJK' in f.name:
                    plt.rcParams['font.family'] = f.name
                    break
            else:
                plt.rcParams['font.family'] = 'sans-serif'
                plt.rcParams['font.sans-serif'] = ['Noto Sans CJK KR', 'Noto Sans CJK', 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False
        plt.rcParams['font.size'] = 12

        charts_dir = os.path.join(project_dir, 'charts')
        os.makedirs(charts_dir, exist_ok=True)

        # 차트 1: TAM/SAM/SOM
        fig, ax = plt.subplots(figsize=(10, 6))
        categories = ['SOM\n(초기 타깃)', 'SAM\n(국내 실질 시장)', 'TAM\n(글로벌 시장)']
        values = [50, 5000, 80000]  # 기본값 (억 원)
        colors = ['#E74C3C', '#F39C12', '#3498DB']
        bars = ax.barh(categories, values, color=colors, height=0.5)
        ax.set_xscale('log')
        ax.set_xlabel('시장 규모 (억 원, 로그 스케일)')
        ax.set_title('시장 규모 계층 구조 (TAM / SAM / SOM)', fontsize=14, fontweight='bold')
        for bar, val in zip(bars, values):
            label = f'{val/10000:.1f}조' if val >= 10000 else f'{val}억'
            ax.text(bar.get_width() * 1.2, bar.get_y() + bar.get_height()/2,
                    f'{label} 원', va='center', fontsize=11, fontweight='bold')
        fig.tight_layout()
        plt.savefig(f'{charts_dir}/market_size.png', dpi=200, bbox_inches='tight', facecolor='white')
        plt.close()

        # 차트 2: 경쟁사 포지셔닝 맵
        fig, ax = plt.subplots(figsize=(10, 7))
        ax.axhspan(5.5, 10.5, xmin=0.5, xmax=1, alpha=0.1, color='#27AE60')
        ax.scatter(9.5, 9.5, s=400, c='#E74C3C', marker='*', zorder=10)
        ax.annotate('우리 제품', (9.5, 9.5), textcoords="offset points",
                    xytext=(10, 8), fontsize=12, fontweight='bold', color='#E74C3C')
        ax.set_xlabel('자동화 수준 →', fontsize=12, fontweight='bold')
        ax.set_ylabel('산업 특화도 ↑', fontsize=12, fontweight='bold')
        ax.set_title('경쟁사 포지셔닝 맵', fontsize=14, fontweight='bold')
        ax.set_xlim(1, 11)
        ax.set_ylim(1, 11)
        ax.grid(True, alpha=0.2)
        ax.axhline(y=5.5, color='gray', ls='--', alpha=0.3)
        ax.axvline(x=5.5, color='gray', ls='--', alpha=0.3)
        fig.tight_layout()
        plt.savefig(f'{charts_dir}/competitor_map.png', dpi=200, bbox_inches='tight', facecolor='white')
        plt.close()

    except Exception as e:
        print(f"차트 생성 실패: {e}")
